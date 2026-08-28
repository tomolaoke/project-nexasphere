"""Multi-file, multi-format ingestion for the "Analyze My Business" workspace.

Design principle, carried over from the rest of NexaSphere: be honest about
what can actually be read. Files are sorted into two buckets, and the
distinction matters:

  DATA    -- structured tables (CSV/XLSX/JSON/TSV) that deterministic
             analytics can compute on. These drive KPIs and findings.
  CONTEXT -- prose extracted from documents (PDF/DOCX/TXT/MD). This is
             business context (targets, policies, notes). It is NEVER
             treated as a measured business figure: a PDF saying "our
             revenue target is 10 million" establishes a target, not
             revenue. Nothing in CONTEXT feeds the analytics engine.

Deliberately NOT supported (rather than faked):
  - Images/OCR: reliable OCR needs a system Tesseract binary, which is not
    available on the free Streamlit Community Cloud runtime. Claiming image
    support that silently returns garbage is worse than declining.
  - Video: there is no reliable, free, in-browser-runtime path from a video
    file to trustworthy business figures. Marked unsupported.

All parsers here are free and open-source (pandas, openpyxl, pypdf,
python-docx). No paid API, no network call, no cloud service -- uploaded
files are parsed in-process and never leave the runtime.
"""
from __future__ import annotations

import io
import json
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional

import pandas as pd

# Extensions we can genuinely parse, by bucket.
STRUCTURED_EXTENSIONS = {".csv", ".tsv", ".xlsx", ".xls", ".json"}
DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
# Recognised-but-declined, so we can explain rather than silently ignore.
UNSUPPORTED_EXTENSIONS = {
    ".png": "image", ".jpg": "image", ".jpeg": "image", ".webp": "image", ".gif": "image",
    ".mp4": "video", ".mov": "video", ".avi": "video", ".mkv": "video", ".webm": "video",
}

MAX_DOCUMENT_CHARS = 20_000  # per document, keeps memory and prompt size sane


@dataclass
class IngestedFile:
    name: str
    kind: str            # "data" | "context" | "unsupported"
    extension: str
    status: str           # "ready" | "extracted" | "declined" | "error"
    detail: str = ""
    frame: Optional[pd.DataFrame] = None   # populated for kind == "data"
    text: str = ""                          # populated for kind == "context"

    @property
    def rows(self) -> Optional[int]:
        return None if self.frame is None else len(self.frame)

    @property
    def columns(self) -> Optional[int]:
        return None if self.frame is None else len(self.frame.columns)

    def summary_row(self) -> dict[str, Any]:
        if self.kind == "data" and self.frame is not None:
            size = f"{len(self.frame):,} rows"
        elif self.kind == "context":
            size = f"{len(self.text):,} chars"
        else:
            size = "—"
        return {
            "File": self.name,
            "Type": self.extension.lstrip(".").upper(),
            "Role": {"data": "Data", "context": "Context", "unsupported": "—"}[self.kind],
            "Size": size,
            "Status": self.status.capitalize(),
            "Detail": self.detail,
        }


# ---------------------------------------------------------------------------
# Structured parsers
# ---------------------------------------------------------------------------

def _read_structured(name: str, ext: str, buf: bytes) -> tuple[Optional[pd.DataFrame], str]:
    """Returns (frame, detail). Raises nothing -- errors come back as detail."""
    try:
        if ext == ".csv":
            return pd.read_csv(io.BytesIO(buf)), ""
        if ext == ".tsv":
            return pd.read_csv(io.BytesIO(buf), sep="\t"), ""
        if ext in (".xlsx", ".xls"):
            # Only the first sheet: silently concatenating sheets with different
            # schemas would fabricate a table the user never had.
            excel = pd.ExcelFile(io.BytesIO(buf))
            first = excel.sheet_names[0]
            detail = ""
            if len(excel.sheet_names) > 1:
                detail = f"Read sheet '{first}' of {len(excel.sheet_names)}; other sheets ignored."
            return excel.parse(first), detail
        if ext == ".json":
            payload = json.loads(buf.decode("utf-8", errors="replace"))
            if isinstance(payload, list):
                return pd.json_normalize(payload), ""
            if isinstance(payload, dict):
                # Prefer the first list-of-records value, else flatten one level.
                for key, value in payload.items():
                    if isinstance(value, list) and value and isinstance(value[0], dict):
                        return pd.json_normalize(value), f"Used array under key '{key}'."
                return pd.json_normalize([payload]), "Flattened a single JSON object."
            return None, "JSON did not contain a table-like structure."
    except Exception as exc:
        return None, f"Could not parse: {exc}"
    return None, "Unrecognised structured format."


# ---------------------------------------------------------------------------
# Document parsers (context only -- never analytics input)
# ---------------------------------------------------------------------------

def _read_document(ext: str, buf: bytes) -> tuple[str, str]:
    """Returns (text, detail)."""
    try:
        if ext in (".txt", ".md"):
            return buf.decode("utf-8", errors="replace")[:MAX_DOCUMENT_CHARS], ""
        if ext == ".pdf":
            try:
                from pypdf import PdfReader
            except ImportError:
                return "", "PDF support unavailable (pypdf not installed)."
            reader = PdfReader(io.BytesIO(buf))
            pages = [(p.extract_text() or "") for p in reader.pages]
            text = "\n".join(pages)[:MAX_DOCUMENT_CHARS]
            if not text.strip():
                return "", ("No selectable text found — this looks like a scanned PDF. "
                             "NexaSphere does not run OCR, so its contents can't be read.")
            return text, f"Extracted text from {len(reader.pages)} page(s)."
        if ext == ".docx":
            try:
                import docx
            except ImportError:
                return "", "DOCX support unavailable (python-docx not installed)."
            document = docx.Document(io.BytesIO(buf))
            text = "\n".join(p.text for p in document.paragraphs)[:MAX_DOCUMENT_CHARS]
            return text, f"Extracted {len(document.paragraphs)} paragraph(s)."
    except Exception as exc:
        return "", f"Could not read: {exc}"
    return "", "Unrecognised document format."


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def ingest_files(uploaded_files) -> list[IngestedFile]:
    """uploaded_files: iterable of Streamlit UploadedFile (or any object with
    .name and .getvalue()/.read()).
    """
    results: list[IngestedFile] = []
    for item in uploaded_files:
        name = getattr(item, "name", "unnamed")
        ext = Path(name).suffix.lower()
        try:
            buf = item.getvalue() if hasattr(item, "getvalue") else item.read()
        except Exception as exc:
            results.append(IngestedFile(name, "unsupported", ext, "error", f"Could not read file: {exc}"))
            continue

        if ext in UNSUPPORTED_EXTENSIONS:
            kind_word = UNSUPPORTED_EXTENSIONS[ext]
            reason = (
                "Image files aren't analyzed — NexaSphere doesn't run OCR, and guessing "
                "at numbers from a picture would break the guarantee that every figure is "
                "traceable to a real calculation."
                if kind_word == "image" else
                "Video files aren't analyzed. There's no reliable free way to turn a video "
                "into trustworthy business figures, so NexaSphere declines rather than "
                "inventing them."
            )
            results.append(IngestedFile(name, "unsupported", ext, "declined", reason))
            continue

        if ext in STRUCTURED_EXTENSIONS:
            frame, detail = _read_structured(name, ext, buf)
            if frame is None or frame.empty:
                results.append(IngestedFile(
                    name, "unsupported", ext, "error",
                    detail or "File contained no data rows.",
                ))
            else:
                frame.columns = [str(c).strip() for c in frame.columns]
                results.append(IngestedFile(name, "data", ext, "ready", detail, frame=frame))
            continue

        if ext in DOCUMENT_EXTENSIONS:
            text, detail = _read_document(ext, buf)
            if not text.strip():
                results.append(IngestedFile(name, "unsupported", ext, "error", detail))
            else:
                results.append(IngestedFile(name, "context", ext, "extracted", detail, text=text))
            continue

        results.append(IngestedFile(
            name, "unsupported", ext, "declined",
            f"'{ext or 'no extension'}' isn't a format NexaSphere reads.",
        ))
    return results


def data_files(files: list[IngestedFile]) -> list[IngestedFile]:
    return [f for f in files if f.kind == "data" and f.frame is not None]


def context_files(files: list[IngestedFile]) -> list[IngestedFile]:
    return [f for f in files if f.kind == "context"]


def detect_relationships(files: list[IngestedFile]) -> list[dict[str, str]]:
    """Finds columns that appear in more than one uploaded table and look like
    identifiers -- a *candidate* join key, surfaced for the user to confirm.

    Deliberately does NOT perform the join. Two files sharing a column named
    "id" or "name" very often are not related at all, and auto-joining them
    would silently produce fabricated combined figures.
    """
    seen: dict[str, list[str]] = {}
    for f in data_files(files):
        for col in f.frame.columns:
            key = str(col).strip().lower()
            # Only consider identifier-shaped column names.
            if key.endswith("_id") or key.endswith("id") or key.endswith("_code"):
                seen.setdefault(key, []).append(f.name)

    relationships = []
    for key, owners in seen.items():
        if len(owners) > 1:
            relationships.append({
                "column": key,
                "files": ", ".join(owners),
                "note": "Shared identifier — these tables may be related.",
            })
    return relationships


def choose_primary_frame(files: list[IngestedFile]) -> Optional[IngestedFile]:
    """The table analytics will run on.

    Chosen by analytical usefulness, not raw size. Picking the largest file
    (the previous rule) meant that uploading a whole business folder selected
    something like a daily inventory snapshot -- hundreds of thousands of rows
    but no revenue column -- over the sales table that actually answers
    business questions. A money column is what unlocks nearly every analysis,
    so it dominates the score; a date column comes next because it unlocks all
    trend and growth work. Row count only breaks ties.

    The choice is always surfaced to the user, and never merges files.
    """
    candidates = data_files(files)
    if not candidates:
        return None

    money_words = ("revenue", "sales", "amount", "total", "price", "profit", "cost", "turnover")
    date_words = ("date", "time", "day", "month", "period")

    def score(f: IngestedFile) -> tuple[int, int, int]:
        cols = [str(c).lower() for c in f.frame.columns]
        has_money = any(
            any(w in c for w in money_words) and pd.api.types.is_numeric_dtype(f.frame[c0])
            for c, c0 in zip(cols, f.frame.columns)
        )
        has_date = any(any(w in c for w in date_words) for c in cols)
        return (int(has_money), int(has_date), len(f.frame))

    return max(candidates, key=score)
